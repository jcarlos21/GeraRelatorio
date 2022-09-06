from docx import Document
fn='REPORT_2022.2-BR38.docx'
document = Document(fn)

my_chapters = ['1   Certificação', '1.1 Metodologia', '1.2 Diagnóstico', '1.3 Status', '2   Resultados', '3   Conclusão']

pn=1    
import re

def tableOfContents(dado):
    # fn=doc
    # document = Document(fn)

    pn=1
    lista = list()

    for p in document.paragraphs:
        r=re.match(dado, p.text)
        if r:
            # print(r.group(),pn)
            lista.append(pn)
        for run in p.runs:
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                pn+=1
                # print('!!','='*50,pn)
    return lista[-1]

print(tableOfContents(my_chapters[5]))
