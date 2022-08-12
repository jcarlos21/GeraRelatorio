from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns

document = Document()

class SetupPage:

    def __init__(self, document):
        self.document = document
    
    def marginsPage(self, top, left, bottom, right):
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(top)
            section.bottom_margin = Cm(bottom)
            section.left_margin = Cm(left)
            section.right_margin = Cm(right)
    

    
    # def create_element(self, name):
    #     return OxmlElement(name)


    # def create_attribute(self, element, name, value):
    #     element.set(ns.qn(name), value)


    # def add_page_number(self, paragraph):
    #     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #     page_run = paragraph.add_run()
    #     t1 = SetupPage(document).create_element('w:t')
    #     SetupPage(document).create_attribute(t1, 'xml:space', 'preserve')
    #     t1.text = 'Page '
    #     page_run._r.append(t1)

    #     page_num_run = paragraph.add_run()

    #     fldChar1 = SetupPage(document).create_element('w:fldChar')
    #     SetupPage(document).create_attribute(fldChar1, 'w:fldCharType', 'begin')

    #     instrText = SetupPage(document).create_element('w:instrText')
    #     SetupPage(document).create_attribute(instrText, 'xml:space', 'preserve')
    #     instrText.text = "PAGE"

    #     fldChar2 = SetupPage(document).create_element('w:fldChar')
    #     SetupPage(document).create_attribute(fldChar2, 'w:fldCharType', 'end')

    #     page_num_run._r.append(fldChar1)
    #     page_num_run._r.append(instrText)
    #     page_num_run._r.append(fldChar2)

    #     of_run = paragraph.add_run()
    #     t2 = SetupPage(document).create_element('w:t')
    #     SetupPage(document).create_attribute(t2, 'xml:space', 'preserve')
    #     t2.text = ' of '
    #     of_run._r.append(t2)

    #     fldChar3 = SetupPage(document).create_element('w:fldChar')
    #     SetupPage(document).create_attribute(fldChar3, 'w:fldCharType', 'begin')

    #     instrText2 = SetupPage(document).create_element('w:instrText')
    #     SetupPage(document).create_attribute(instrText2, 'xml:space', 'preserve')
    #     instrText2.text = "NUMPAGES"

    #     fldChar4 = SetupPage(document).create_element('w:fldChar')
    #     SetupPage(document).create_attribute(fldChar4, 'w:fldCharType', 'end')

    #     num_pages_run = paragraph.add_run()
    #     num_pages_run._r.append(fldChar3)
    #     num_pages_run._r.append(instrText2)
    #     num_pages_run._r.append(fldChar4)
