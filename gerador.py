from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from conteudos import CriaTexto
from stylesTexts import StylesText
from docx.shared import Inches
from pageSetup import SetupPage
from analysisFunctions import AnalysisFunc
from dictFill import WriteDict
from docx.oxml import OxmlElement, ns

import datetime

global d, m, y, data, my_chapters, data_up
d = datetime.datetime.today().strftime('%d')
m = datetime.datetime.today().strftime('%B')
y = datetime.datetime.today().strftime('%Y')


class TextGenerator:

    def __init__(self, document='', ticket='', data=f'{d} de {m} de {y}', cause_correction='', range_test=30, technician='',
                    technician_reg=1, scholarship='', scholarship_reg=1, observations='', dataDict={'dado': 1}):
        
        self.document = document
        self.ticket = ticket
        self.data = data
        self.cause_correction = cause_correction
        self.range_test = range_test
        self.technician = technician
        self.technician_reg = technician_reg
        self.scholarship = scholarship
        self.scholarship_reg = scholarship_reg
        self.observations = observations
        self.dataDict = dataDict
    
    def text_genetator(self):

        texto = CriaTexto(self.document)
        estilos = StylesText(self.document)
        pageConfig = SetupPage(self.document)
        analise = AnalysisFunc(self.document)
        
        pageConfig.marginsPage(3.0, 3.0, 2.0, 2.0)

        # Capa
        # Formatação 1:
        fte = ('Arial', 1, False, False, 12, False)

        texto.textoSimples(self.technician.strip(), fte[0], fte[1], fte[2], fte[3], fte[4], fte[5])
        texto.textoSimples(f'Matrícula: {self.technician_reg.strip()}', fte[0], fte[1], fte[2], fte[3], fte[4], fte[5])
        texto.textoSimples(f'{self.scholarship.strip()}', fte[0], fte[1], fte[2], fte[3], fte[4], fte[5])
        texto.textoSimples(f'Matrícula: {self.scholarship_reg.strip()}', fte[0], fte[1], fte[2], fte[3], fte[4], fte[5])
        texto.addNewLine(6)

        titulo = f'Rede Giga Metrópole\nRelatório de Conformidade Referente ao Bilhete {self.ticket}'
        texto.textoSimples(titulo, 'Arial', 1, True, False, 12, False)
        texto.addNewLine(6)

        cabecalho = """Ponto de Presença da Rede Nacional de Ensino e Pesquisa no Rio Grande do Norte - POP-RN\nRede GigaMetropole\nSetor de Infraestrutura"""
        texto.textoSimples(cabecalho, fte[0], fte[1], fte[2], fte[3], fte[4], fte[5])
        texto.addNewLine(5)

        texto.textoSimples('Natal - RN', fte[0], fte[1], fte[2], fte[3], fte[4], fte[5])
        texto.textoSimples(self.data.strip(), fte[0], fte[1], fte[2], fte[3], fte[4], fte[5])

        # Sumário

        self.document.add_page_break()  # Quebra 1

        texto.textoSimples('Sumário', 'Arial', 1, True, False, 12, False)
        texto.addNewLine(0)
        base = ['Texto_1', 'Texto_2', 'Texto_3', 'Texto_4', 'Texto_5', 'Texto_6']
        texto.sumario_inicial(base)

        # Número de página

        section = self.document.sections[0]
        footer = section.footer
        texto.addNovaSection()
        pageConfig.add_page_number(footer.paragraphs[0])

        # Manutenção Corretiva RGM

        self.document.add_page_break()  # Quebra 2

        texto.textoSimples ('Manutenção Corretiva RGM', 'Arial', 1, True, False, 12, False)
        texto.addNewLine(0)

        p = self.document.add_paragraph()

        t1 = p.add_run('Objetivo: certificar o serviço de manutenção corretiva realizado pela empresa Interjato Soluções (')
        estilos.addStyles(t1, 'Arial', False, False, 12)
        estilos.addStyles(p.add_run(f'bilhete {self.ticket.strip()}'), 'Arial', True, False, 12)

        t3 = p.add_run(') para restabelecer à conectividade GPON na(s) célula(s) ')
        estilos.addStyles(t3, 'Arial', False, False, 12)
        texto.repeteListaEmUmaLinha(list(self.dataDict.keys()), p, 'Arial', True, False, 12)

        t5 = p.add_run('. Os dados apresentados nesse documento foram obtidos a partir do monitoramento da rede GPON realizado pelo software ')
        estilos.addStyles(t5, 'Arial', False, False, 12)
        estilos.addStyles(p.add_run('Grafana.'), 'Arial', True, True, 12)
        texto.addNewLine(0)

        texto.textoSimples ('Entidade(s) afetada(s) pelo rompimento do cabo de fibras óptica:', 'Arial', 3, False, False, 12, False)
        texto.addNewLine(0)

        texto.repetMarcadoreEntidade(self.dataDict, 'Arial', 0, negrito=True, italico=False, tam=12)
        texto.addNewLine(0)

        texto.textoSimples('Local da Ocorrência:', 'Arial', 3, False, False, 12, False)
        texto.addNewLine(0)

        texto.imprimeMarcadorEndereco(self.dataDict, 'Arial', negrito=False, italico=False, tam=12)

        q = self.document.add_paragraph()  # necessário pois, ao usar o método .textoSimples(), um novo .add_paragraph() (ou seja, um novo parágrafo) é iniciado.
        pf = q.paragraph_format
        pf.left_indent = Inches(0.5)
        q.style = 'List Bullet'
        texto.repeteListaEmUmaLinha('Trecho(s): ', q, fonte='Arial', negrito=False, italico=False, tam=12)

        texto.imprimeCaixaEntidade(self.dataDict, q, fonte='Arial', negrito=True, italico=False, tam=12)
        texto.addNewLine(0)

        texto.textoSimples('Informações do Cabo:', 'Arial', 3, False, False, 12, False)
        texto.addNewLine(0)
        texto.addMarcadores(self.cause_correction, 'Arial', 0, False, False, 12)

        texto.textoFormat(p, 3, 1.50, 0, 0)
        texto.textoFormat(q, 0, 1.50, 0, 0)

        # Página Simples 1

        self.document.add_page_break()  # Quebra 3

        texto1 = """Todos os ativos GPON da Rede Gigametrópole são monitorados pelo software GRAFANA. Dentre os parâmetros monitorados, são de interesse nesse processo de certificação os valores de potência óptica recebidos que são enviados periodicamente pelas ONU(s). A certificação é baseadas nos seguintes requisitos:"""

        texto.textoSimples('1   Certificação', 'Arial', 3, True, False, 12, False)
        texto.addNewLine(0)
        texto.textoSimples('1.1 Metodologia', 'Arial', 3, True, False, 12, False)
        texto.addNewLine(0)

        texto.textoSimples(texto1, 'Arial', 3, False, False, 12, True)
        texto.addNewLine(0)

        requisitos = ['Comparação entre os valores de potência recebidos antes e depois do incidente;',
        'Comparação entre os valores de potência recebido em cada cliente (ONU) afetado pelo incidente e a média de potência recebida nos outros clientes da mesma célula;',
        'Análise do comportamento do sinal recebido na(s) ONU(s), buscando identificar oscilações relevantes (maiores que 1 dB entorno do valor médio) na magnitude do sinal.']

        texto.repetMarcadores(requisitos, 'Arial', 3, False, False, 12)
        texto.addNewLine(0)

        texto.textoSimples('1.2 Diagnóstico', 'Arial', 3, True, False, 12, False)
        texto.addNewLine(0)
        texto2 = """A comparação dos resultados obtidos pelo monitoramento apresentados no(s) gráfico(s) da(s) figura(s) apresentada(s) no Resultados e nos dados da tabela 3, mostram que os níveis de potência óptica recebidos na(s) ONU(s) são coerentes. São sintetizados nas tabelas 1 e 2, as respostas aos requisitos estabelecidos e o diagnóstico da manutenção corretiva. """
        texto.textoSimples(texto2, 'Arial', 3, False, False, 12, True)
        texto.addNewLine(0)

        # Requisitos de teste

        r1 = 'R1 – Os valores de potência permaneceram na mesma ordem de grandeza antes e depois do incidente?'
        texto.addMarcadores(r1, 'Arial', 3, False, False, 12)

        pj = self.document.add_paragraph()
        pk = pj.paragraph_format
        pk.left_indent = Inches(0.5)
        r = pj.add_run('R2 – Considerando que o valor médio de potência óptica recebida nas ONU(s) das escolas, nas células ')
        pj.style = 'List Bullet'
        estilos.addStyles(r, 'Arial', False, False, 12)
        texto.textoFormat(pj, 3, 1.50, 0, 0)

        texto.repeteListaEmUmaLinha(list(self.dataDict.keys()), pj, 'Arial', True, False, 12)

        r2b = ' é de '
        texto.repeteListaEmUmaLinha(r2b, pj, 'Arial', False, False, 12)

        for caixa in self.dataDict.keys():  # Lembre-se de criar um método para tratar os loops
            for escola in self.dataDict[caixa].keys():
                texto.repeteListaEmUmaLinha(str(self.dataDict[caixa][escola][1]) + ' dBm, ', pj, 'Arial', False, False, 12)

        r2c = ' a potência obtida na(s) ONU(s) após o reparo estão na mesma ordem de grandeza do valor médio?'
        texto.repeteListaEmUmaLinha(r2c, pj, 'Arial', False, False, 12)

        r3 = 'R3 – A oscilação no sinal recebido é aceitável?'
        texto.addMarcadores(r3, 'Arial', 3, False, False, 12)

        # Página Simples 2

        self.document.add_page_break()  # Quebra 4
        
        texto.textoSimples('Legendas das respostas aos requisitos:', 'Arial', 3, False, False, 12, False)
        texto.addNewLine(0)
        texto.textoSimples('1.  OK – Em conformidade;', 'Arial', 3, False, False, 12, True)
        texto.textoSimples('2.  X – Não atende ao requisito.', 'Arial', 3, False, False, 12, True)
        texto.addNewLine(0)


        # Tabela 1

        texto.textoSimples('Tabela 1 – Resultado do diagnóstico', fte[0], fte[1], fte[2], fte[3], fte[4], fte[5])

        table = self.document.add_table(rows=1, cols=4)
        row = table.rows[0].cells
        headerRow = ['ESCOLA', 'R1', 'R2', 'R3']
        texto.alimentaTabela(row, headerRow, 'Arial', 12)

        dataDiagnostic = list()
        dataDiagnosticStatus = list()
        dataDiagnosticPot = list()

        for enS in self.dataDict.keys():
            for enI in self.dataDict[enS].keys():
                a1 = analise.rR1(self.dataDict[enS][enI][2], self.dataDict[enS][enI][3])
                a2 = analise.rR2(self.dataDict[enS][enI][1], self.dataDict[enS][enI][3])
                a3 = analise.rR3(self.dataDict[enS][enI][3])
                status = analise.rStatus(a1, a2, a3)
                dataDiagnostic.append([enI, a1, a2, a3])
                dataDiagnosticStatus.append([enI, status])
                dataDiagnosticPot.append([enI, str(self.dataDict[enS][enI][2]), str(self.dataDict[enS][enI][3])])

        for escola, r1, r2, r3 in dataDiagnostic:  # Adding a row and then adding data in it.
            row = table.add_row().cells
            listaLinhas = [escola, r1, r2, r3]
            texto.alimentaTabela(row, listaLinhas, 'Arial', 12)

        for cell1 in table.columns[0].cells:
            cell1.width = Inches(2.8)

        for i in range(1, 4):
            for cell in table.columns[i].cells:
                cell.width = Inches(0.5)

        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        texto.addNewLine(0)

        # Tabela 2

        texto.textoSimples('1.3 Status', 'Arial', 3, True, False, 12, False)
        texto.addNewLine(0)
        texto.textoSimples('O serviço de manutenção corretiva é qualificado conforme os status apresentados na tabela 2.', 'Arial', 3, False, False, 12, True)
        texto.addNewLine(0)

        texto.textoSimples('Tabela 2 – Status do serviço de manutenção corretiva', fte[0], fte[1], fte[2], fte[3], fte[4], fte[5])

        table2 = self.document.add_table(rows=1, cols=2)
        row2 = table2.rows[0].cells
        headerRow2 = ['PONTO ATENDIDO', 'STATUS']
        texto.alimentaTabela(row2, headerRow2, 'Arial', 12)

        for escola, status in dataDiagnosticStatus:  # Adding a row and then adding data in it.
            row2 = table2.add_row().cells
            listaLinhas2 = [escola, status]
            texto.alimentaTabela(row2, listaLinhas2, 'Arial', 12)

        for cell2 in table2.columns[0].cells:
            cell2.width = Inches(2.8)

        for i in range(1, 2):
            for cell in table2.columns[i].cells:
                cell.width = Inches(1.5)

        table2.style = 'Table Grid'
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        texto.addNewLine(0)

        texto.textoSimples('2   Resultados', 'Arial', 3, True, False, 12, False)
        texto.addNewLine(0)

        # Tabela 3

        texto.textoSimples('Tabela 3 – Valor médio da potência óptica recebida nas ONUs', fte[0], fte[1], fte[2], fte[3], fte[4], fte[5])

        table3 = self.document.add_table(rows=1, cols=3)
        row3 = table3.rows[0].cells
        headerRow3 = ['ESCOLA', 'PRxA[dBm]', 'PRxB[dBm]']
        texto.alimentaTabela(row3, headerRow3, 'Arial', 12)

        for escola, ptA, ptD in dataDiagnosticPot:  # Adding a row and then adding data in it.
            row3 = table3.add_row().cells
            listaLinhas3 = [escola, ptA, ptD]
            texto.alimentaTabela(row3, listaLinhas3, 'Arial', 12)

        for cell3 in table3.columns[0].cells:
            cell3.width = Inches(2.8)

        for i in range(1, 3):
            for cell in table3.columns[i].cells:
                cell.width = Inches(1.0)

        table3.style = 'Table Grid'
        table3.alignment = WD_TABLE_ALIGNMENT.CENTER
        texto.addNewLine(0)

        texto.addMarcadores ('PRxA [dBm] - Potência óptica recebida na ONU antes do incidente;', 'Arial', 3, False, False, 12)
        texto.addMarcadores ('PRxB [dBm] - Potência óptica recebida na ONU após o reparo.', 'Arial', 3, False, False, 12)

        # Página Simples 3

        self.document.add_page_break()  # Quebra 4

        texto3 = f'No(s) gráfico(s) apresentado(s) na(s) figura(s) a seguir, os resultados mostram o comportamento do sinal recebido durante o período de {self.range_test}, considerando antes e após o serviço de reparação ser executado. É importante ressaltar que no decorrer do período de amostragem apresentado no(s) gráfico(s) podem ocorrer intervalos sem amostras, como o período de observação é grande e os dados são enviados pelas ONU, é possível que em algum momento o equipamento seja desligado.'
        texto.textoSimples(texto3, 'Arial', 3, False, False, 12, True)
        texto.addNewLine(0)

        for i in range(0, len(dataDiagnosticStatus)):
            texto.textoSimples(f'Figura {i+1} - Monitoramento GPON: potência óptica recebida na ONU da {dataDiagnosticStatus[i][0]}', fte[0], fte[1], fte[2], fte[3], fte[4], fte[5])
            self.document.add_picture(f'img/0{str(i+1)}.JPG')
            texto.addNewLine(0)
        
        # Página Simples 4

        self.document.add_page_break()  # Quebra 6

        texto.textoSimples('3   Conclusão', 'Arial', 3, True, False, 12, False)
        texto.addNewLine(0)
        texto4 = f'{self.observations}. Diante do dados analisados, conclui-se que os resultados apresentados nesse documento certificam que o serviço de manutenção corretiva foi executado em conformidade com os padrões exigidos e sanando todas as pendências, garantindo o correto funcionamento da rede.'
        texto.textoSimples(texto4, 'Arial', 3, False, False, 12, True)

    def generator_docx(self):

        self.document.save(f"output_docx/REPORT_{self.ticket}.docx")

        texto = CriaTexto(self.document)

        base = ['Texto_1', 'Texto_2', 'Texto_3', 'Texto_4', 'Texto_5', 'Texto_6']
        my_chapters_search = ['1   Certificação', '1.1 Metodologia', '1.2 Diagnóstico', '1.3 Status', '2   Resultados', '3   Conclusão']
        my_chapters_write = ['1    Certificação...', '1.1 Metodologia..', '1.2 Diagnóstico...', '1.3 Status...........', '2    Resultados....', '3    Conclusão.....']
        texto.sumario (f"output_docx/REPORT_{self.ticket}.docx", base, my_chapters_search, my_chapters_write)

        return 'Documento gerado com sucesso!'
