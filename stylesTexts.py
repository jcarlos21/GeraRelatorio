from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Pt

document = Document()

class StylesText:

    def __init__(self, document):
        self.document = document
        pass

    def addStyles(self, alinhamento, dado, fonte, negrito, italico, tamanhoFonte):
        texto = self.document.add_paragraph()
        texto.alignment = alinhamento
        texto.paragraph_format.line_spacing = 1.50
        texto.paragraph_format.space_before = Pt(0)
        texto.paragraph_format.space_after = Pt(0)
        r = texto.add_run(dado)
        r.font.name = fonte
        r.font.size = Pt(tamanhoFonte)
        r.font.bold = negrito
        r.font.italic = italico