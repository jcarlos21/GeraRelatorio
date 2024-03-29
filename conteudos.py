from re import template
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from stylesTexts import StylesText
from docx.shared import Inches

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tableofcontents as tb

document = Document()
estilos = StylesText(document)


class CriaTexto:

    def __init__(self, document):
        self.document = document
    
    def addNewLine(self, qtd):
        line = self.document.add_paragraph("\n"*qtd)
        line.paragraph_format.line_spacing = 1.50
        line.paragraph_format.space_before = Pt(0)
        line.paragraph_format.space_after = Pt(0)
    
    def textoSimples (self, texto, fonte, alinhamento, negrito, italico, tam, identacao):
        """
            O alinhamento pode ser 1, 2, 3 e 4
                0 - LEFT: Left-aligned
                1 - CENTER: Center-aligned
                2 - RIGHT: Right-aligned
                3 - JUSTIFY: Fully justified.
        """

        paragrafo = self.document.add_paragraph()
        
        if identacao:
            pf = paragrafo.paragraph_format
            pf.first_line_indent = Inches(0.5)

        CriaTexto(document).textoFormat(paragrafo, alinhamento, 1.50, 0, 0)

        r = paragrafo.add_run(texto)

        estilos.addStyles(r, fonte, negrito, italico, tam)

    def addMarcadores (self, dado, fonte, alinhamento, negrito, italico, tam):
        marcador = self.document.add_paragraph()

        paragraph_format = marcador.paragraph_format
        paragraph_format.left_indent = Inches(0.5)

        CriaTexto(self.document).textoFormat(marcador, alinhamento, 1.50, 0, 0)

        marcador.style = 'List Bullet'

        r = marcador.add_run(dado)
        estilos.addStyles(r, fonte, negrito, italico, tam)
    
    def textoFormat(self, instancia, alinhamento, space, space_after, space_before):
        instancia.alignment = alinhamento
        instancia.paragraph_format.line_spacing = space
        instancia.paragraph_format.space_before = Pt(space_after)
        instancia.paragraph_format.space_after = Pt(space_before)

    def alimentaTabela (self, rowTable, listRows, fonte, tam):
        for i in range(0, len(listRows)):
            rowTable[i].text = listRows[i]
            p = rowTable[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.runs[0]
            estilos.addStyles(r, fonte, False, False, tam)
    
    def addNovaSection(self):
        self.document.add_section(WD_SECTION.CONTINUOUS)
        self.document.sections[1].footer.is_linked_to_previous = False

    def repeteListaEmUmaLinha(self, dado, p, fonte, negrito, italico, tam):
        
        if type(dado) == list:
            for i in range(0, len(dado)):
                if i < len(dado) - 1:
                    t = p.add_run(f'{dado[i]}, ')
                    estilos.addStyles(t, fonte, negrito, italico, tam)
                else:
                    t = p.add_run(f'{dado[i]}')
                    estilos.addStyles(t, fonte, negrito, italico, tam)
        else:
            t = p.add_run(dado)
            estilos.addStyles(t, fonte, negrito, italico, tam)
        
    def repetMarcadores(self, dado, fonte, alinhamento, negrito, italico, tam):
        for i in range(0, len(dado)):
            CriaTexto(self.document).addMarcadores (dado[i], fonte, alinhamento, negrito, italico, tam)
    
    def repetMarcadoreEntidade(self, dado, fonte, alinhamento, negrito, italico, tam):
        for entidadeSuperior in dado.keys():
            for entidadeInferior in dado[entidadeSuperior].keys():
                CriaTexto(self.document).addMarcadores (entidadeInferior, fonte, alinhamento, negrito, italico, tam)

    def imprimeMarcadorEndereco(self, dado, fonte, negrito, italico, tam):
        for entidadeSuiperior in dado.keys():
            for entidadeInferior in dado[entidadeSuiperior].keys():
                CriaTexto(self.document).addMarcadores(f'Endereço: {dado[entidadeSuiperior][entidadeInferior][0]}', fonte, 0, negrito, italico, tam)
    
    def imprimeCaixaEntidade(self, dado, p, fonte, negrito, italico, tam):
        for entidadeSuperior in dado.keys():
            for entidadeInferior in dado[entidadeSuperior].keys():
                estilos.addStyles(p.add_run(f'{entidadeSuperior.upper()} - {entidadeInferior.upper()}; '), fonte, negrito, italico, tam)
    
    def sumario_inicial(self, my_chapters):
        for chapter in my_chapters:
            CriaTexto(self.document).textoSimples(chapter, 'Arial', 3, False, False, 12, False)
    
    def sumario(self, name_doc, text_start, my_chapters_s, my_chapters_w):

        template_document = Document(name_doc)

        for i in range(0, len(text_start)):
            for paragraph in template_document.paragraphs:
                num = tb.tableOfContents(my_chapters_s[i], name_doc)
                tb.replace_text_in_paragraph(paragraph, text_start[i], my_chapters_w[i] + 110*'.' + f' {num}')
        template_document.save(name_doc)
